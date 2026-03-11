"""Google Drive Service fuer Kandidaten-Upload an Marketing-Agentur.

Erstellt Ordner-Struktur: {PLZ}_{Primary_Role}/
Laedt hoch: Profil.pdf + Transkript.docx
"""

import io
import logging
from uuid import UUID

from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

from app.config import settings

logger = logging.getLogger(__name__)

# Google OAuth Token-Endpoint
TOKEN_URI = "https://oauth2.googleapis.com/token"


class GoogleDriveService:
    """Upload-Service fuer Google Drive (OAuth 2.0)."""

    def __init__(self):
        self._service = None

    @property
    def is_available(self) -> bool:
        """Prueft ob alle Google Drive Credentials konfiguriert sind."""
        return bool(
            settings.google_drive_client_id
            and settings.google_drive_client_secret
            and settings.google_drive_refresh_token
            and settings.google_drive_folder_id
        )

    def _get_service(self):
        """Erstellt authentifizierten Google Drive API Service (cached)."""
        if self._service:
            return self._service

        creds = Credentials(
            token=None,
            refresh_token=settings.google_drive_refresh_token,
            token_uri=TOKEN_URI,
            client_id=settings.google_drive_client_id,
            client_secret=settings.google_drive_client_secret,
            scopes=["https://www.googleapis.com/auth/drive.file"],
        )

        self._service = build("drive", "v3", credentials=creds)
        return self._service

    def _find_or_create_folder(self, folder_name: str, parent_id: str) -> str:
        """Findet existierenden Ordner oder erstellt neuen.

        Returns: Folder ID
        """
        service = self._get_service()

        # Suche nach existierendem Ordner
        query = (
            f"name = '{folder_name}' "
            f"and '{parent_id}' in parents "
            f"and mimeType = 'application/vnd.google-apps.folder' "
            f"and trashed = false"
        )
        results = service.files().list(q=query, fields="files(id, name)").execute()
        files = results.get("files", [])

        if files:
            logger.info(f"Drive-Ordner gefunden: {folder_name} ({files[0]['id']})")
            return files[0]["id"]

        # Neuen Ordner erstellen
        metadata = {
            "name": folder_name,
            "mimeType": "application/vnd.google-apps.folder",
            "parents": [parent_id],
        }
        folder = service.files().create(body=metadata, fields="id").execute()
        folder_id = folder["id"]
        logger.info(f"Drive-Ordner erstellt: {folder_name} ({folder_id})")
        return folder_id

    def _upload_file(
        self, file_name: str, file_bytes: bytes, mime_type: str, folder_id: str
    ) -> str:
        """Laedt Datei in einen Drive-Ordner hoch.

        Wenn Datei mit gleichem Namen existiert, wird sie ueberschrieben.
        Returns: File ID
        """
        service = self._get_service()

        # Pruefen ob Datei bereits existiert
        query = (
            f"name = '{file_name}' "
            f"and '{folder_id}' in parents "
            f"and trashed = false"
        )
        results = service.files().list(q=query, fields="files(id)").execute()
        existing = results.get("files", [])

        media = MediaIoBaseUpload(
            io.BytesIO(file_bytes), mimetype=mime_type, resumable=True
        )

        if existing:
            # Update bestehende Datei
            file_id = existing[0]["id"]
            service.files().update(fileId=file_id, media_body=media).execute()
            logger.info(f"Drive-Datei aktualisiert: {file_name} ({file_id})")
            return file_id

        # Neue Datei erstellen
        metadata = {"name": file_name, "parents": [folder_id]}
        result = (
            service.files().create(body=metadata, media_body=media, fields="id").execute()
        )
        file_id = result["id"]
        logger.info(f"Drive-Datei hochgeladen: {file_name} ({file_id})")
        return file_id

    async def upload_candidate_to_drive(
        self,
        candidate_id: UUID,
        postal_code: str,
        primary_role: str,
        pdf_bytes: bytes,
        docx_bytes: bytes | None = None,
    ) -> dict:
        """Laedt Kandidaten-Profil + Transkript in Google Drive hoch.

        Ordner-Struktur: Kandidaten/{PLZ}_{Primary_Role}/
        Dateien: Profil.pdf, Transkript.docx

        Returns: Dict mit folder_id, pdf_file_id, docx_file_id, folder_name
        """
        import asyncio

        if not self.is_available:
            raise RuntimeError("Google Drive ist nicht konfiguriert")

        # Ordnername: z.B. "34434_Bilanzbuchhalterin"
        plz = (postal_code or "00000").strip()
        role = (primary_role or "Unbekannt").strip().replace("/", "-")
        folder_name = f"{plz}_{role}"

        root_folder_id = settings.google_drive_folder_id

        # Blocking API-Calls in Executor ausfuehren
        loop = asyncio.get_event_loop()

        def _sync_upload():
            # Unterordner erstellen/finden
            folder_id = self._find_or_create_folder(folder_name, root_folder_id)

            result = {
                "folder_id": folder_id,
                "folder_name": folder_name,
                "pdf_file_id": None,
                "docx_file_id": None,
            }

            # Profil-PDF hochladen
            result["pdf_file_id"] = self._upload_file(
                "Profil.pdf", pdf_bytes, "application/pdf", folder_id
            )

            # Transkript.docx hochladen (wenn vorhanden)
            if docx_bytes:
                result["docx_file_id"] = self._upload_file(
                    "Transkript.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    folder_id,
                )

            return result

        result = await loop.run_in_executor(None, _sync_upload)
        logger.info(
            f"Kandidat {candidate_id} an Marketing gesendet: "
            f"Ordner={result['folder_name']}, PDF={result['pdf_file_id']}, "
            f"DOCX={result['docx_file_id']}"
        )
        return result


def create_transcript_docx(transcript_text: str, candidate_name: str = "") -> bytes:
    """Erstellt ein Word-Dokument aus dem Transkript-Text.

    Returns: DOCX als bytes
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Seitenraender
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titel
    title = doc.add_heading("Gespraechstranskript", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if candidate_name:
        subtitle = doc.add_paragraph()
        run = subtitle.add_run(f"Kandidat: {candidate_name}")
        run.font.size = Pt(12)
        run.bold = True

    doc.add_paragraph("")  # Leerzeile

    # Transkript-Text in Absaetze aufteilen
    paragraphs = transcript_text.strip().split("\n")
    for para_text in paragraphs:
        stripped = para_text.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()
        # Sprecher-Erkennung (z.B. "Recruiter:" oder "Kandidat:")
        if ":" in stripped and len(stripped.split(":")[0]) < 30:
            speaker, _, rest = stripped.partition(":")
            run_speaker = p.add_run(f"{speaker}:")
            run_speaker.bold = True
            run_speaker.font.size = Pt(11)
            if rest.strip():
                run_text = p.add_run(f" {rest.strip()}")
                run_text.font.size = Pt(11)
        else:
            run = p.add_run(stripped)
            run.font.size = Pt(11)

    # In Bytes umwandeln
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
